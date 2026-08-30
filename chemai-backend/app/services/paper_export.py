# -*- coding: utf-8 -*-
"""试卷导出服务。

将 Paper 渲染为打印友好 HTML（零依赖）或可选 Word .docx（python-docx）。
导出内容：标题、按 sort_order 排序的题目列表、每题分值、总分；
含答案/解析开关（不含答案时同时隐藏解析）。
"""
from __future__ import annotations

import html as _html
from io import BytesIO
from typing import Optional

from app.core.enums import QuestionType
from app.models import Question

# 题型中文展示映射（导出文档展示用）。
_TYPE_LABELS: dict[QuestionType, str] = {
    QuestionType.SINGLE_CHOICE: "单选题",
    QuestionType.MULTI_CHOICE: "多选题",
    QuestionType.TRUE_FALSE: "判断题",
    QuestionType.FILL_BLANK: "填空题",
    QuestionType.SHORT_ANSWER: "简答题",
    QuestionType.ESSAY: "论述题",
    QuestionType.CALCULATION: "计算题",
    QuestionType.EXPERIMENT: "实验题",
    QuestionType.INFERENCE: "推断题",
}


def total_score(questions: list[Question]) -> float:
    """试卷总分（由题目分值求和）。"""
    return sum(q.score or 0.0 for q in questions)


def _type_label(q: Question) -> str:
    return _TYPE_LABELS.get(q.type, q.type.value)


def render_html(
    title: str,
    duration: Optional[int],
    questions: list[Question],
    include_answer: bool = True,
    include_analysis: bool = True,
) -> str:
    """渲染打印友好 HTML。

    不含答案时同时隐藏解析（解析会泄露答案推理）。
    """
    show_analysis = include_answer and include_analysis
    total = total_score(questions)

    body: list[str] = []
    body.append(f"<h1>{_html.escape(title)}</h1>")
    if duration:
        body.append(f"<p class='meta'>时长：{duration} 分钟</p>")
    body.append(f"<p class='meta'>总分：{total} 分</p>")

    for i, q in enumerate(questions, 1):
        body.append(
            f"<div class='q'>"
            f"<div class='q-head'>"
            f"<span class='q-no'>{i}.</span> "
            f"<span class='q-type'>[{_type_label(q)}]</span> "
            f"<span class='q-score'>({q.score} 分)</span>"
            f"</div>"
            f"<div class='q-content'>{_html.escape(q.content or '')}</div>"
        )
        if q.options:
            opts = "".join(
                f"<div class='opt'>{chr(65 + j)}. {_html.escape(str(o))}</div>"
                for j, o in enumerate(q.options)
            )
            body.append(f"<div class='opts'>{opts}</div>")
        if include_answer:
            body.append(f"<div class='answer'>答案：{_html.escape(q.answer or '')}</div>")
        if show_analysis and q.analysis:
            body.append(f"<div class='analysis'>解析：{_html.escape(q.analysis)}</div>")
        body.append("</div>")

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>{_html.escape(title)}</title>
<style>
  body {{ font-family: -apple-system, "PingFang SC", "Microsoft YaHei", sans-serif; line-height: 1.6; max-width: 720px; margin: 24px auto; padding: 0 16px; color: #222; }}
  h1 {{ font-size: 22px; text-align: center; }}
  .meta {{ text-align: center; color: #666; }}
  .q {{ margin: 18px 0; page-break-inside: avoid; }}
  .q-head {{ font-weight: 600; }}
  .q-type {{ color: #555; }}
  .q-score {{ color: #888; }}
  .opt {{ margin-left: 20px; }}
  .answer {{ color: #1a6b3c; margin-top: 6px; }}
  .analysis {{ color: #666; margin-top: 4px; }}
  @media print {{ .answer {{ color: #000; }} }}
</style>
</head>
<body>
{chr(10).join(body)}
</body>
</html>"""


def render_docx(
    title: str,
    duration: Optional[int],
    questions: list[Question],
    include_answer: bool = True,
    include_analysis: bool = True,
) -> bytes:
    """渲染 Word .docx（python-docx），返回文件字节。"""
    from docx import Document

    show_analysis = include_answer and include_analysis
    total = total_score(questions)

    doc = Document()
    doc.add_heading(title, level=1)
    if duration:
        doc.add_paragraph(f"时长：{duration} 分钟")
    doc.add_paragraph(f"总分：{total} 分")

    for i, q in enumerate(questions, 1):
        doc.add_paragraph(f"{i}. （{_type_label(q)}，{q.score} 分）{q.content or ''}")
        if q.options:
            for j, opt in enumerate(q.options):
                doc.add_paragraph(f"    {chr(65 + j)}. {opt}")
        if include_answer:
            doc.add_paragraph(f"答案：{q.answer or ''}")
        if show_analysis and q.analysis:
            doc.add_paragraph(f"解析：{q.analysis}")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
