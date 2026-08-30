# -*- coding: utf-8 -*-
"""试卷导出服务测试（L1）：HTML/docx 内容、答案/解析开关、总分与顺序。"""
from __future__ import annotations

from io import BytesIO
from types import SimpleNamespace

from docx import Document

from app.core.enums import QuestionType
from app.services.paper_export import render_docx, render_html, total_score


def _q(
    content: str,
    type_: QuestionType = QuestionType.SINGLE_CHOICE,
    options: list[str] | None = None,
    answer: str = "A",
    analysis: str = "因为……所以选 A",
    score: float = 2.0,
) -> SimpleNamespace:
    return SimpleNamespace(
        content=content,
        type=type_,
        options=options,
        answer=answer,
        analysis=analysis,
        score=score,
    )


class TestRenderHtml:
    def test_html_contains_title_content_score(self) -> None:
        qs = [_q("下列反应中属于氧化还原反应的是？", options=["A", "B"])]
        html = render_html("期中测试", 60, qs)
        assert "期中测试" in html
        assert "总分：2.0 分" in html
        assert "下列反应中属于氧化还原反应的是？" in html
        assert "A." in html
        assert "单选题" in html
        assert "答案：A" in html
        assert "解析：因为……所以选 A" in html

    def test_html_hides_answer_and_analysis_when_disabled(self) -> None:
        html = render_html("期中测试", None, [_q("题目")], include_answer=False)
        assert "答案" not in html
        assert "解析" not in html

    def test_html_analysis_switch_independent(self) -> None:
        # 含答案但关闭解析：答案显示，解析隐藏
        html = render_html("期中测试", None, [_q("题目")], include_answer=True, include_analysis=False)
        assert "答案：A" in html
        assert "解析" not in html

    def test_total_score_sums(self) -> None:
        qs = [_q("a", score=1.5), _q("b", score=3.0)]
        assert total_score(qs) == 4.5


class TestRenderDocx:
    def test_docx_returns_parseable_bytes(self) -> None:
        data = render_docx("期中测试", 60, [_q("下列反应中属于氧化还原反应的是？", options=["A", "B"])])
        assert isinstance(data, bytes) and data
        doc = Document(BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "期中测试" in text
        assert "下列反应中属于氧化还原反应的是？" in text
        assert "答案：A" in text

    def test_docx_hides_answer_when_disabled(self) -> None:
        data = render_docx("期中测试", None, [_q("题目")], include_answer=False)
        doc = Document(BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "答案" not in text
        assert "解析" not in text
